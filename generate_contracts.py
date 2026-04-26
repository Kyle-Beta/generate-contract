#!/usr/bin/env python3
"""
合同批量生成脚本
用法：python generate_contracts.py [选项]

数据源：Excel 文件，第一行为字段名（列头），后续每行生成一份合同
模板：Word (.docx) 文件，用 {{字段名}} 作为占位符

Python >= 3.7，支持 .xlsx / .xlsm / .xls
"""
from __future__ import annotations  # 让 3.7-3.9 也能用 list[dict]、str | None 语法

import argparse
import datetime
import os
import re
import sys
from pathlib import Path

import openpyxl
from docx import Document


# ── 占位符正则，匹配 {{任意内容}} ──────────────────────────────────────
PLACEHOLDER_RE = re.compile(r"\{\{(.+?)\}\}")


# ── 核心替换：合并 run 后再替换 ──────────────────────────────────────────
def _replace_in_paragraph(paragraph, mapping: dict):
    """
    将段落内所有 {{key}} 替换为 mapping[key]。
    Word 常把一个占位符拆分到多个 run，需先合并再替换。
    """
    full_text = "".join(run.text for run in paragraph.runs)
    if "{{" not in full_text:
        return

    # 替换占位符
    def replacer(m):
        key = m.group(1).strip()
        return str(mapping.get(key, m.group(0)))  # 找不到则保留原文

    new_text = PLACEHOLDER_RE.sub(replacer, full_text)
    if new_text == full_text:
        return

    # 把替换后的文本写回第一个 run，其余 run 清空（保留格式）
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def _replace_in_table(table, mapping: dict):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, mapping)
            # 递归嵌套表格
            for nested in cell.tables:
                _replace_in_table(nested, mapping)


def _replace_in_section(section_part, mapping: dict):
    """替换页眉/页脚中的占位符"""
    if section_part is None:
        return
    for paragraph in section_part.paragraphs:
        _replace_in_paragraph(paragraph, mapping)
    for table in section_part.tables:
        _replace_in_table(table, mapping)


def fill_template(template_path: str, mapping: dict, output_path: str):
    """根据 mapping 填充模板，保存到 output_path"""
    doc = Document(template_path)

    # 正文段落
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, mapping)

    # 正文表格
    for table in doc.tables:
        _replace_in_table(table, mapping)

    # 页眉页脚
    for section in doc.sections:
        _replace_in_section(section.header, mapping)
        _replace_in_section(section.footer, mapping)
        _replace_in_section(section.even_page_header, mapping)
        _replace_in_section(section.even_page_footer, mapping)
        _replace_in_section(section.first_page_header, mapping)
        _replace_in_section(section.first_page_footer, mapping)

    doc.save(output_path)


# ── 读取 Excel ────────────────────────────────────────────────────────
def _cell_to_str(value) -> str:
    """将 Excel 单元格值统一转为字符串，处理日期、浮点整数等情况"""
    if value is None:
        return ""
    if isinstance(value, (datetime.datetime, datetime.date)):
        return value.strftime("%Y年%m月%d日")
    if isinstance(value, float):
        # 整数型浮点（如 150000.0）去掉小数点
        return str(int(value)) if value == int(value) else str(value)
    if isinstance(value, bool):
        return "是" if value else "否"
    return str(value)


def _read_xlsx(excel_path: str) -> list[dict]:
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        raise ValueError("Excel 文件为空")
    headers = [
        str(h).strip() if h is not None else f"列{i+1}"
        for i, h in enumerate(rows[0])
    ]
    records = []
    for row in rows[1:]:
        if all(v is None for v in row):
            continue
        records.append({h: _cell_to_str(v) for h, v in zip(headers, row)})
    return records


def _read_xls(excel_path: str) -> list[dict]:
    """读取旧版 .xls 文件，需要 xlrd 库"""
    try:
        import xlrd
    except ImportError:
        raise ImportError(
            "读取 .xls 文件需要安装 xlrd：pip install xlrd"
        )
    wb = xlrd.open_workbook(excel_path)
    ws = wb.sheet_by_index(0)
    if ws.nrows == 0:
        raise ValueError("Excel 文件为空")

    headers = [
        str(ws.cell_value(0, c)).strip() or f"列{c+1}"
        for c in range(ws.ncols)
    ]
    records = []
    for r in range(1, ws.nrows):
        row_vals = []
        for c in range(ws.ncols):
            cell = ws.cell(r, c)
            # xlrd 日期类型需要单独转换
            if cell.ctype == xlrd.XL_CELL_DATE:
                dt = xlrd.xldate_as_datetime(cell.value, wb.datemode)
                row_vals.append(dt.strftime("%Y年%m月%d日"))
            elif cell.ctype == xlrd.XL_CELL_EMPTY:
                row_vals.append("")
            elif cell.ctype == xlrd.XL_CELL_NUMBER:
                v = cell.value
                row_vals.append(str(int(v)) if v == int(v) else str(v))
            elif cell.ctype == xlrd.XL_CELL_BOOLEAN:
                row_vals.append("是" if cell.value else "否")
            else:
                row_vals.append(str(cell.value))
        if all(v == "" for v in row_vals):
            continue
        records.append(dict(zip(headers, row_vals)))
    return records


def read_excel(excel_path: str) -> list[dict]:
    """
    读取 Excel，返回 list[dict]。
    第一行为列头（字段名），后续每行为一条记录。
    支持 .xlsx / .xlsm（openpyxl）和 .xls（xlrd）。
    """
    ext = Path(excel_path).suffix.lower()
    if ext == ".xls":
        return _read_xls(excel_path)
    elif ext in (".xlsx", ".xlsm"):
        return _read_xlsx(excel_path)
    else:
        raise ValueError(f"不支持的文件格式：{ext}，请使用 .xlsx 或 .xls")


def list_excel_fields(excel_path: str) -> tuple[list[str], int]:
    """返回 Excel 字段名和记录数。"""
    records = read_excel(excel_path)
    if not records:
        return [], 0
    return list(records[0].keys()), len(records)


# ── 生成输出文件名 ────────────────────────────────────────────────────
def make_output_filename(record: dict, name_field: str | None,
                         index: int, output_dir: str) -> str:
    """
    优先用 name_field 对应的值作为文件名；
    否则用 合同_001.docx 格式。
    """
    if name_field and name_field in record and record[name_field]:
        # 去除文件名非法字符
        safe = re.sub(r'[\\/:*?"<>|]', "_", record[name_field])
        filename = f"{safe}.docx"
    else:
        filename = f"合同_{index:03d}.docx"
    return os.path.join(output_dir, filename)


def _dedupe_output_path(output_path: str, used_paths: set[str]) -> str:
    """避免同名文件被覆盖。"""
    candidate = Path(output_path)
    stem = candidate.stem
    suffix = candidate.suffix
    parent = candidate.parent
    counter = 2

    while str(candidate) in used_paths or candidate.exists():
        candidate = parent / f"{stem}_{counter}{suffix}"
        counter += 1
    return str(candidate)


def generate_contracts(
    excel_path: str,
    template_path: str,
    output_dir: str = "output_contracts",
    name_field: str | None = None,
) -> dict:
    """
    批量生成合同并返回结果摘要。

    返回：
        {
            "records": int,
            "success": int,
            "failed": int,
            "output_dir": str,
            "generated_files": list[str],
            "errors": list[str],
        }
    """
    records = read_excel(excel_path)
    if not records:
        raise ValueError("Excel 中没有数据行（第一行之后无内容）")

    os.makedirs(output_dir, exist_ok=True)

    success = 0
    failed = 0
    generated_files: list[str] = []
    errors: list[str] = []
    used_paths: set[str] = set()

    for i, record in enumerate(records, start=1):
        output_path = make_output_filename(record, name_field, i, output_dir)
        output_path = _dedupe_output_path(output_path, used_paths)
        used_paths.add(output_path)
        try:
            fill_template(template_path, record, output_path)
            generated_files.append(output_path)
            success += 1
        except Exception as e:
            errors.append(f"[{i:03d}] {e}")
            failed += 1

    return {
        "records": len(records),
        "success": success,
        "failed": failed,
        "output_dir": os.path.abspath(output_dir),
        "generated_files": generated_files,
        "errors": errors,
    }


# ── 主流程 ────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(
        description="根据 Excel 数据源和 Word 模板批量生成合同"
    )
    parser.add_argument("excel", help="Excel 数据源文件路径（.xlsx）")
    parser.add_argument("template", help="Word 合同模板路径（.docx）")
    parser.add_argument(
        "-o", "--output-dir",
        default="output_contracts",
        help="输出目录（默认：output_contracts）"
    )
    parser.add_argument(
        "-n", "--name-field",
        default=None,
        help="用于命名输出文件的字段名（如：乙方名称）；不填则用序号"
    )
    parser.add_argument(
        "--list-fields",
        action="store_true",
        help="仅列出 Excel 中的字段名，不生成合同"
    )
    args = parser.parse_args()

    # 检查输入文件
    for path, label in [(args.excel, "Excel"), (args.template, "Word 模板")]:
        if not os.path.exists(path):
            print(f"错误：{label} 文件不存在：{path}", file=sys.stderr)
            sys.exit(1)

    if args.list_fields:
        print(f"读取数据：{args.excel}")
        fields, count = list_excel_fields(args.excel)
        if not fields:
            print("Excel 中没有数据行（第一行之后无内容）")
            sys.exit(0)
        print(f"\nExcel 字段名（共 {len(fields)} 个）：")
        for field in fields:
            print(f"  {{{{ {field} }}}}")
        print(f"\n共 {count} 条记录")
        return

    print(f"读取数据：{args.excel}")
    print(f"模板：{args.template}")
    print(f"输出目录：{args.output_dir}")
    result = generate_contracts(
        args.excel,
        args.template,
        output_dir=args.output_dir,
        name_field=args.name_field,
    )

    print(f"共 {result['records']} 条记录，开始生成...\n")
    for i, file_path in enumerate(result["generated_files"], start=1):
        print(f"  [{i:03d}] ✓ {os.path.basename(file_path)}")
    for error in result["errors"]:
        print(f"  ✗ 生成失败：{error}", file=sys.stderr)

    print(f"\n完成：成功 {result['success']} 份，失败 {result['failed']} 份")
    print(f"文件保存在：{result['output_dir']}/")


def _fix_windows_encoding():
    """Windows 终端默认 GBK，强制 UTF-8 输出避免中文乱码"""
    if sys.platform != "win32":
        return
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    except AttributeError:
        pass  # Python < 3.7 不支持 reconfigure


if __name__ == "__main__":
    _fix_windows_encoding()
    main()
